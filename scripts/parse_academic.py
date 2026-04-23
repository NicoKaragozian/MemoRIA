import hashlib
import json
import logging
import re
import statistics
from pathlib import Path

logger = logging.getLogger(__name__)

_REF_HEADERS = re.compile(
    r'^\s*(References|Bibliography|Bibliografía|Referencias|Works Cited|Fuentes)\s*$',
    re.IGNORECASE,
)
_PAGE_NUM = re.compile(r'^\s*\d{1,4}\s*$')

# Separa oraciones de forma simple (no exhaustiva, suficiente para estadísticas)
_SENT_SPLIT = re.compile(r'(?<=[.!?])\s+')


def _extract_pdf(filepath: str) -> tuple[list[str], int]:
    import pdfplumber
    pages = []
    with pdfplumber.open(filepath) as pdf:
        total = len(pdf.pages)
        for page in pdf.pages:
            raw = page.extract_text() or ""
            pages.append(_clean_page(raw))
    return pages, total


def _extract_docx(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _clean_page(text: str) -> str:
    lines = text.splitlines()
    clean = []
    for line in lines:
        if _PAGE_NUM.match(line):
            continue
        if _REF_HEADERS.match(line):
            break
        clean.append(line)
    return "\n".join(clean)


def _detect_repeated_lines(page_texts: list[str], threshold: float = 0.3) -> set[str]:
    from collections import Counter
    n_pages = len(page_texts)
    if n_pages < 3:
        return set()
    line_counts: Counter = Counter()
    for page_text in page_texts:
        for line in set(page_text.splitlines()):
            stripped = line.strip()
            if stripped:
                line_counts[stripped] += 1
    return {line for line, cnt in line_counts.items() if cnt / n_pages >= threshold}


def _chunk_hash(text: str) -> str:
    normalized = re.sub(r'\s+', ' ', text.lower().strip())
    return hashlib.sha1(normalized.encode()).hexdigest()


def _median_sentence_len(text: str) -> float:
    sentences = [s.strip() for s in _SENT_SPLIT.split(text) if s.strip()]
    if not sentences:
        return 0.0
    lengths = [len(s.split()) for s in sentences]
    return statistics.median(lengths)


def split_into_paragraph_chunks(
    text: str,
    min_words: int = 300,
    max_words: int = 600,
) -> list[str]:
    """
    Divide el texto en chunks respetando límites de párrafo.
    Concatena párrafos hasta alcanzar min_words sin superar max_words.
    Si un párrafo solo supera max_words, se divide en mitad de oraciones.
    """
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    chunks = []
    current_parts: list[str] = []
    current_words = 0

    for para in paragraphs:
        para_words = len(para.split())
        if current_words + para_words > max_words and current_parts:
            chunks.append("\n\n".join(current_parts))
            current_parts = []
            current_words = 0
        current_parts.append(para)
        current_words += para_words
        if current_words >= min_words:
            chunks.append("\n\n".join(current_parts))
            current_parts = []
            current_words = 0

    if current_parts and current_words >= 100:
        chunks.append("\n\n".join(current_parts))

    return chunks


def parse_academic_folder(folder: str, min_chunk_len: int = 300) -> list[dict]:
    """
    Procesa todos los PDFs y DOCX en la carpeta.
    Elimina headers/footers repetidos, referencias y chunks duplicados entre archivos.
    No aplica anonimización — datos son uso local exclusivo.
    """
    folder_path = Path(folder)
    files = list(folder_path.glob("**/*.pdf")) + list(folder_path.glob("**/*.docx"))

    examples = []
    seen_hashes: set[str] = set()
    failed = 0

    for filepath in files:
        logger.info("Procesando: %s", filepath.name)
        try:
            if filepath.suffix == ".pdf":
                page_texts, _ = _extract_pdf(str(filepath))
                repeated = _detect_repeated_lines(page_texts)
                cleaned_pages = []
                for pt in page_texts:
                    lines = [l for l in pt.splitlines() if l.strip() not in repeated]
                    cleaned_pages.append("\n".join(lines))
                text = "\n\n".join(cleaned_pages)
            else:
                text = _extract_docx(str(filepath))

            if not text or len(text) < 500:
                logger.warning("  Texto muy corto en %s, saltando", filepath.name)
                continue

            chunks = split_into_paragraph_chunks(text)
            added = 0
            for chunk in chunks:
                if len(chunk) < min_chunk_len:
                    continue
                # Descartar chunks con oraciones muy cortas (listas, TOC, bullets)
                if _median_sentence_len(chunk) < 15:
                    continue
                h = _chunk_hash(chunk)
                if h in seen_hashes:
                    continue
                seen_hashes.add(h)
                examples.append({
                    "text": chunk,
                    "source": filepath.name,
                    "register": "academic",
                })
                added += 1
            logger.info("  → %d chunks (de %d candidatos)", added, len(chunks))

        except Exception as e:
            logger.error("  Error en %s: %s", filepath.name, e)
            failed += 1

    if failed:
        logger.warning("%d/%d archivos fallaron", failed, len(files))

    return examples


if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)
    folder = sys.argv[1] if len(sys.argv) > 1 else "data/raw/academic"
    examples = parse_academic_folder(folder)
    print(f"\nTotal: {len(examples)} chunks académicos")
    output = "academic_parsed.jsonl"
    with open(output, "w", encoding="utf-8") as f:
        for ex in examples:
            f.write(json.dumps(ex, ensure_ascii=False) + "\n")
    print(f"Guardado en {output}")
